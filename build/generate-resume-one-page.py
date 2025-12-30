from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

def create_resume():
    doc = Document()

    # VERY narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Smaller default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(9.5)

    # --- Header ---
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = header.add_run('TAHA MERGHANI')
    name_run.bold = True
    name_run.font.size = Pt(14)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run('taha.y.merghani@gmail.com | ')
    add_hyperlink(contact, 'linkedin.com/in/tahamerghani', 'https://linkedin.com/in/tahamerghani')
    contact.add_run(' | ')
    add_hyperlink(contact, 'github.com/taha-y-merghani', 'https://github.com/taha-y-merghani')
    contact.add_run(' | ')
    add_hyperlink(contact, 'taha-y-merghani.github.io', 'https://taha-y-merghani.github.io')

    header.paragraph_format.space_after = Pt(0)
    contact.paragraph_format.space_after = Pt(4)

    # --- Summary ---
    h = doc.add_heading('SUMMARY', level=1)
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(2)
    summary = doc.add_paragraph(
        "AI Research Engineer focused on LLM robustness and evaluation. Publications at NeurIPS/NAACL, 4.0 GPA, built Siri tooling at Apple still in production 8 years later. Combining linguistic expertise with modern engineering (Ollama, Whisper)."
    )
    summary.paragraph_format.space_after = Pt(4)

    # --- Experience ---
    h = doc.add_heading('EXPERIENCE', level=1)
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(2)

    # Independent
    p = doc.add_paragraph()
    p.add_run('Independent AI Research Engineer').bold = True
    p.add_run(' | April 2023 – Present')
    p.paragraph_format.space_after = Pt(1)
    bullets = [
        "Authored technical analysis on LLM limitations achieving >200k views, recognition from Google Chief Scientist Jeff Dean",
        "Built local voice interfaces using Whisper, Ollama, Mistral 7B optimizing for latency/privacy on consumer hardware"
    ]
    for b in bullets:
        bp = doc.add_paragraph(b, style='List Bullet')
        bp.paragraph_format.space_after = Pt(0)

    # Mesa
    p = doc.add_paragraph()
    p.add_run('Solutions Architect, Mesa Associates').bold = True
    p.add_run(' | Nov 2022 – Apr 2023')
    p.paragraph_format.space_after = Pt(1)
    bp = doc.add_paragraph("Deployed computer vision models on Oracle Cloud for automated electrical equipment inspection", style='List Bullet')
    bp.paragraph_format.space_after = Pt(0)

    # Decooda
    p = doc.add_paragraph()
    p.add_run('AI Software Engineer, Decooda').bold = True
    p.add_run(' | Jul 2019 – Aug 2020')
    p.paragraph_format.space_after = Pt(1)
    bp = doc.add_paragraph("Built ETL pipelines using Matillion and Snowflake for customer analytics and ML model training", style='List Bullet')
    bp.paragraph_format.space_after = Pt(0)

    # Georgia Tech
    p = doc.add_paragraph()
    p.add_run('Graduate Research Assistant, Georgia Institute of Technology').bold = True
    p.add_run(' | Aug 2017 – May 2019')
    p.paragraph_format.space_after = Pt(1)

    # NeurIPS with hyperlink
    bp = doc.add_paragraph(style='List Bullet')
    bp.add_run('Published ')
    add_hyperlink(bp, 'Geolinguistic Analysis via Twitter (NeurIPS 2018 Workshop)', 'https://taha-y-merghani.github.io/lyon-geolinguistic-variation-2018.pdf')
    bp.add_run(' and ')
    add_hyperlink(bp, 'NAACL 2018 Workshop', 'https://aclanthology.org/W18-1602/')
    bp.paragraph_format.space_after = Pt(0)

    bp = doc.add_paragraph("Research on conversational AI under Dr. Mark Riedl", style='List Bullet')
    bp.paragraph_format.space_after = Pt(0)

    # Apple
    p = doc.add_paragraph()
    p.add_run('Software Engineering Intern, Apple').bold = True
    p.add_run(' | Summer 2016')
    p.paragraph_format.space_after = Pt(1)
    bp = doc.add_paragraph(style='List Bullet')
    bp.add_run('Built debugging and visualization tools for Siri NLU pipeline, ')
    bp.add_run('still in production').bold = True
    bp.add_run(' use after 8 years')
    bp.paragraph_format.space_after = Pt(0)

    # MIT
    p = doc.add_paragraph()
    p.add_run('Research Intern, MIT CSAIL').bold = True
    p.add_run(' | Summer 2015')
    p.paragraph_format.space_after = Pt(1)
    bp = doc.add_paragraph("Developed lexical models for Egyptian Arabic ASR under James Glass; grapheme lexicon outperformed diacritized approaches", style='List Bullet')
    bp.paragraph_format.space_after = Pt(0)

    # --- Education ---
    h = doc.add_heading('EDUCATION', level=1)
    h.paragraph_format.space_before = Pt(4)
    h.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    p.add_run('M.S. Computer Science, Georgia Institute of Technology').bold = True
    p.add_run(' | 2017–2019 | GPA: 3.87/4.0')
    p.paragraph_format.space_after = Pt(1)

    p = doc.add_paragraph()
    p.add_run('B.S. Computer Engineering, Jackson State University').bold = True
    p.add_run(' | 2013–2017 | GPA: 4.0/4.0 | Jacksonian Award (Highest GPA)')
    p.paragraph_format.space_after = Pt(4)

    # --- Publications ---
    h = doc.add_heading('PUBLICATIONS & WRITING', level=1)
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(2)

    pubs = [
        ("From Sudan to Silicon Valley: Beyond the Resume (Shared by Jeff Dean)", "https://medium.com/@tahaymerghani/from-sudan-to-silicon-valley-beyond-the-resume-a2e6bd3eedb4"),
        ("What 5,000 Hours of Mastering Tekken Taught Me About AI Research (#1 on r/ArtificialIntelligence)", "https://medium.com/@tahaymerghani/a-machine-learning-researcher-spent-close-to-5-000-hours-on-tekken-and-reached-top-0-5-a42c96877214"),
        ("Balusu, Merghani, Eisenstein. Stylistic Variation in Social Media POS Tagging. NAACL Workshop 2018", "https://aclanthology.org/W18-1602/")
    ]

    for text, url in pubs:
        p = doc.add_paragraph(style='List Bullet')
        add_hyperlink(p, text, url)
        p.paragraph_format.space_after = Pt(0)

    # --- Skills ---
    h = doc.add_heading('SKILLS', level=1)
    h.paragraph_format.space_before = Pt(4)
    h.paragraph_format.space_after = Pt(2)
    p = doc.add_paragraph("Python, PyTorch, Hugging Face, LLMs (Ollama, Whisper), SQL, Snowflake, AWS, Docker, TypeScript | Arabic (Native), English (Near-native)")
    p.paragraph_format.space_after = Pt(0)

    doc.save('Taha_Merghani_Resume.docx')
    print("✓ ONE-PAGE resume generated successfully")

create_resume()
