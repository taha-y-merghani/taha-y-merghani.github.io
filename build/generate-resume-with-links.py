from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Hyperlink styling
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

def create_resume():
    doc = Document()

    # Set narrow margins for one-page fit
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # --- Styles ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)  # Reduced from 11 for space

    # --- Header ---
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = header.add_run('TAHA MERGHANI')
    name_run.bold = True
    name_run.font.size = Pt(16)  # Reduced from 18

    # Contact info with hyperlinks
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run('taha.y.merghani@gmail.com | ')
    add_hyperlink(contact, 'linkedin.com/in/tahamerghani', 'https://linkedin.com/in/tahamerghani')
    contact.add_run(' | ')
    add_hyperlink(contact, 'github.com/taha-y-merghani', 'https://github.com/taha-y-merghani')
    contact.add_run(' | ')
    add_hyperlink(contact, 'taha-y-merghani.github.io', 'https://taha-y-merghani.github.io')

    # Reduce spacing
    header.paragraph_format.space_after = Pt(2)
    contact.paragraph_format.space_after = Pt(6)

    # --- Summary ---
    doc.add_heading('SUMMARY', level=1).paragraph_format.space_after = Pt(4)
    summary_text = doc.add_paragraph(
        "AI Research Engineer with a focus on LLM robustness, evaluation, and linguistic nuance. "
        "Background includes 4.0 GPA, publications at NeurIPS/NAACL, and building Siri debugging tooling "
        "at Apple that remains in production eight years later. Combining deep linguistic expertise with modern "
        "engineering (Ollama, Whisper) to build reliable AI systems."
    )
    summary_text.paragraph_format.space_after = Pt(6)

    # --- Experience ---
    doc.add_heading('EXPERIENCE', level=1).paragraph_format.space_after = Pt(4)

    # Role 1: Independent
    p = doc.add_paragraph()
    p.add_run('Independent AI Research Engineer').bold = True
    p.add_run('\tApril 2023 – Present').bold = True
    p.paragraph_format.space_after = Pt(2)

    bullets_indep = [
        "Authored technical analysis on LLM limitations and AI research culture, achieving >200k views and recognition from Google Chief Scientist Jeff Dean.",
        "Developing local-first voice interfaces using OpenAI Whisper, Ollama, and Mistral 7B to optimize latency and privacy on consumer hardware.",
        "Consulting on data labeling quality and RLHF pipelines for expert networks (NewtonX)."
    ]
    for b in bullets_indep:
        p = doc.add_paragraph(b, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    # Role 2: Mesa
    p = doc.add_paragraph()
    p.add_run('Solutions Architect, Mesa Associates').bold = True
    p.add_run('\tNov 2022 – Apr 2023').bold = True
    p.paragraph_format.space_after = Pt(2)

    bullets_mesa = [
        "Deployed computer vision models on Oracle Cloud for automated electrical equipment inspection.",
        "Built Python tooling for drone image processing and automated reporting workflows."
    ]
    for b in bullets_mesa:
        p = doc.add_paragraph(b, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    # Role 3: Decooda
    p = doc.add_paragraph()
    p.add_run('AI Software Engineer, Decooda').bold = True
    p.add_run('\tJul 2019 – Aug 2020').bold = True
    p.paragraph_format.space_after = Pt(2)

    bullets_decooda = [
        "Built ETL pipelines using Matillion and Snowflake for customer analytics and ML model training."
    ]
    for b in bullets_decooda:
        p = doc.add_paragraph(b, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    # Role 4: Georgia Tech
    p = doc.add_paragraph()
    p.add_run('Graduate Research Assistant, Georgia Institute of Technology').bold = True
    p.add_run('\tAug 2017 – May 2019').bold = True
    p.paragraph_format.space_after = Pt(2)

    bullets_gt = [
        "Published at NeurIPS 2018 Workshop (geolinguistic variation via kernel methods) and NAACL 2018 Workshop.",
        "Conducted research on human-computer interaction for conversational AI under Dr. Mark Riedl."
    ]
    for b in bullets_gt:
        p = doc.add_paragraph(b, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    # Role 5: Apple
    p = doc.add_paragraph()
    p.add_run('Software Engineering Intern, Apple').bold = True
    p.add_run('\tSummer 2016').bold = True
    p.paragraph_format.space_after = Pt(2)

    bullets_apple = [
        "Built internal debugging and visualization tools for the Siri NLU pipeline.",
        "Tools streamlined engineering workflows and remain in production use eight years later."
    ]
    for b in bullets_apple:
        p = doc.add_paragraph(b, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    # Role 6: MIT
    p = doc.add_paragraph()
    p.add_run('Research Intern, MIT CSAIL').bold = True
    p.add_run('\tSummer 2015').bold = True
    p.paragraph_format.space_after = Pt(2)

    bullets_mit = [
        "Developed lexical models for Egyptian Arabic ASR under advisor James Glass.",
        "Demonstrated that grapheme lexicon approaches outperformed complex diacritized approaches for dialectal speech."
    ]
    for b in bullets_mit:
        p = doc.add_paragraph(b, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    # --- Education ---
    doc.add_heading('EDUCATION', level=1).paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.add_run('M.S. Computer Science, Georgia Institute of Technology').bold = True
    p.add_run('\t2017 – 2019')
    p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph('GPA: 3.87/4.0 | Advisors: Jacob Eisenstein, Mark Riedl').paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.add_run('B.S. Computer Engineering, Jackson State University').bold = True
    p.add_run('\t2013 – 2017')
    p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph('GPA: 4.0/4.0 | Jacksonian Award (Highest GPA in College of Science & Engineering)').paragraph_format.space_after = Pt(6)

    # --- Publications & Writing ---
    doc.add_heading('SELECTED WRITING & PUBLICATIONS', level=1).paragraph_format.space_after = Pt(4)

    writing_items = [
        ("From Sudan to Silicon Valley: Beyond the Resume (Shared by Jeff Dean/Google Chief Scientist).", "https://medium.com/@tahaymerghani/from-sudan-to-silicon-valley-beyond-the-resume-a2e6bd3eedb4"),
        ("What 5,000 Hours of Mastering Tekken Taught Me About AI Research (#1 on r/ArtificialIntelligence).", "https://medium.com/@tahaymerghani/a-machine-learning-researcher-spent-close-to-5-000-hours-on-tekken-and-reached-top-0-5-a42c96877214"),
        ("Merghani, Eisenstein. Geolinguistic Analysis via Twitter. NeurIPS Workshop (Black in AI) 2018.", None),
        ("Balusu, Merghani, Eisenstein. Stylistic Variation in Social Media POS Tagging. NAACL Workshop 2018.", "https://aclanthology.org/W18-1602/")
    ]

    for text, url in writing_items:
        p = doc.add_paragraph(style='List Bullet')
        if url:
            add_hyperlink(p, text, url)
        else:
            p.add_run(text)
        p.paragraph_format.space_after = Pt(2)

    # --- Skills ---
    doc.add_heading('SKILLS', level=1).paragraph_format.space_after = Pt(4)
    skills_p = doc.add_paragraph(
        "Technical: Python, PyTorch, Hugging Face, LLMs (Ollama, Whisper), SQL, Snowflake, AWS, Docker, Kernel Methods, TypeScript. "
        "Languages: Arabic (Native), English (Near-native)."
    )
    skills_p.paragraph_format.space_after = Pt(0)

    doc.save('Taha_Merghani_Resume.docx')
    print("✓ Resume with hyperlinks generated successfully.")

create_resume()
