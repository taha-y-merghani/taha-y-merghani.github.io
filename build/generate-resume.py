from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_resume():
    doc = Document()

    # --- Styles ---
    # Set default font to something professional (e.g., Calibri or Arial)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- Header ---
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = header.add_run('TAHA MERGHANI')
    name_run.bold = True
    name_run.font.size = Pt(18)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run('taha.y.merghani@gmail.com | linkedin.com/in/tahamerghani | github.com/taha-y-merghani | taha-y-merghani.github.io')

    doc.add_paragraph() # Spacer

    # --- Summary ---
    doc.add_heading('SUMMARY', level=1)
    summary_text = doc.add_paragraph(
        "AI Research Engineer with a focus on LLM robustness, evaluation, and linguistic nuance. "
        "Background includes 4.0 GPA, publications at NeurIPS/NAACL, and building Siri debugging tooling "
        "at Apple that remains in production eight years later. Combining deep linguistic expertise with modern "
        "engineering (Ollama, Whisper) to build reliable AI systems."
    )

    # --- Experience ---
    doc.add_heading('EXPERIENCE', level=1)

    # Role 1: Independent (New - Fills the Gap)
    p = doc.add_paragraph()
    p.add_run('Independent AI Research Engineer').bold = True
    p.add_run('\tApril 2023 – Present').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    bullets_indep = [
        "Authored technical analysis on LLM limitations and AI research culture, achieving >200k views and recognition from Google Chief Scientist Jeff Dean.",
        "Developing local-first voice interfaces using OpenAI Whisper, Ollama, and Mistral 7B to optimize latency and privacy on consumer hardware.",
        "Consulting on data labeling quality and RLHF pipelines for expert networks (NewtonX)."
    ]
    for b in bullets_indep:
        doc.add_paragraph(b, style='List Bullet')

    # Role 2: Mesa
    p = doc.add_paragraph()
    p.add_run('Solutions Architect, Mesa Associates').bold = True
    p.add_run('\tNov 2022 – Apr 2023').bold = True

    doc.add_paragraph("Remote", style='Normal').italic = True

    bullets_mesa = [
        "Deployed computer vision models on Oracle Cloud for automated electrical equipment inspection.",
        "Built Python tooling for drone image processing and automated reporting workflows."
    ]
    for b in bullets_mesa:
        doc.add_paragraph(b, style='List Bullet')

    # Role 3: Decooda
    p = doc.add_paragraph()
    p.add_run('AI Software Engineer, Decooda').bold = True
    p.add_run('\tJul 2019 – Aug 2020').bold = True

    doc.add_paragraph("Atlanta, GA", style='Normal').italic = True

    bullets_decooda = [
        "Built ETL pipelines using Matillion and Snowflake for customer analytics and ML model training."
    ]
    for b in bullets_decooda:
        doc.add_paragraph(b, style='List Bullet')

    # Role 4: Georgia Tech
    p = doc.add_paragraph()
    p.add_run('Graduate Research Assistant, Georgia Institute of Technology').bold = True
    p.add_run('\tAug 2017 – May 2019').bold = True

    doc.add_paragraph("Atlanta, GA", style='Normal').italic = True

    bullets_gt = [
        "Published at NeurIPS 2018 Workshop (geolinguistic variation via kernel methods) and NAACL 2018 Workshop.",
        "Conducted research on human-computer interaction for conversational AI under Dr. Mark Riedl."
    ]
    for b in bullets_gt:
        doc.add_paragraph(b, style='List Bullet')

    # Role 5: Apple
    p = doc.add_paragraph()
    p.add_run('Software Engineering Intern, Apple').bold = True
    p.add_run('\tSummer 2016').bold = True

    doc.add_paragraph("Cupertino, CA", style='Normal').italic = True

    bullets_apple = [
        "Built internal debugging and visualization tools for the Siri NLU pipeline.",
        "Tools streamlined engineering workflows and remain in production use eight years later."
    ]
    for b in bullets_apple:
        doc.add_paragraph(b, style='List Bullet')

    # Role 6: MIT
    p = doc.add_paragraph()
    p.add_run('Research Intern, MIT CSAIL').bold = True
    p.add_run('\tSummer 2015').bold = True

    doc.add_paragraph("Cambridge, MA", style='Normal').italic = True

    bullets_mit = [
        "Developed lexical models for Egyptian Arabic ASR under advisor James Glass.",
        "Demonstrated that grapheme lexicon approaches outperformed complex diacritized approaches for dialectal speech."
    ]
    for b in bullets_mit:
        doc.add_paragraph(b, style='List Bullet')

    # --- Education ---
    doc.add_heading('EDUCATION', level=1)

    p = doc.add_paragraph()
    p.add_run('M.S. Computer Science, Georgia Institute of Technology').bold = True
    p.add_run('\t2017 – 2019')
    doc.add_paragraph('GPA: 3.87/4.0 | Advisors: Jacob Eisenstein, Mark Riedl')

    p = doc.add_paragraph()
    p.add_run('B.S. Computer Engineering, Jackson State University').bold = True
    p.add_run('\t2013 – 2017')
    doc.add_paragraph('GPA: 4.0/4.0 | Jacksonian Award (Highest GPA in College of Science & Engineering)')

    # --- Publications & Writing ---
    doc.add_heading('SELECTED WRITING & PUBLICATIONS', level=1)

    writing = [
        "From Sudan to Silicon Valley: Beyond the Resume (Shared by Jeff Dean/Google Chief Scientist).",
        "What 5,000 Hours of Mastering Tekken Taught Me About AI Research (#1 on r/ArtificialIntelligence).",
        "Merghani, Eisenstein. Geolinguistic Analysis via Twitter. NeurIPS Workshop (Black in AI) 2018.",
        "Balusu, Merghani, Eisenstein. Stylistic Variation in Social Media POS Tagging. NAACL Workshop 2018."
    ]
    for w in writing:
        doc.add_paragraph(w, style='List Bullet')

    # --- Skills ---
    doc.add_heading('SKILLS', level=1)
    skills = (
        "Technical: Python, PyTorch, Hugging Face, LLMs (Ollama, Whisper), SQL, Snowflake, AWS, Docker, Kernel Methods, TypeScript.\n"
        "Languages: Arabic (Native), English (Near-native)."
    )
    doc.add_paragraph(skills)

    doc.save('Taha_Merghani_Resume.docx')
    print("Resume generated successfully.")

create_resume()
